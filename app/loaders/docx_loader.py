"""
DocxLoader — converts Word documents to Markdown.

Two-stage approach:
  Stage 1 — MarkItDown (fast, offline, zero API calls)
             Handles: text, headings, lists, simple tables, hyperlinks
  Stage 2 — LangChain LLM enrichment (only when complexity detected)
             Triggers: images present, nested tables detected

Complexity detection uses python-docx to inspect the document BEFORE
parsing — so we never pay LLM cost unless actually needed.

LLM enrichment handles:
  - Images: GPT-4o Vision via your LangChain OpenAI wrapper → alt-text descriptions
  - Nested tables: LLM cleans up garbled MarkItDown output

Corporate-safe: no model downloads, no HuggingFace, uses existing
LangChain ChatOpenAI wrapper already approved in your org.
"""

import asyncio
from io import BytesIO

import structlog

from app.loaders.base import BaseLoader, ParsedDocument
from app.connectors.base import RawDocument
from app.core.exceptions import DocumentParsingError

log = structlog.get_logger(__name__)

# Thresholds for complexity detection
_MIN_TEXT_RATIO = 0.05  # if <5% of content is text → likely scanned/image-heavy


class DocxLoader(BaseLoader):

    def __init__(self, llm=None) -> None:
        """
        Args:
            llm: Optional LangChain ChatOpenAI (or any ChatModel) instance.
                 If None, LLM enrichment is disabled — library-only mode.
                 Pass your langchain_openai.ChatOpenAI instance.
        """
        from markitdown import MarkItDown

        self._converter = MarkItDown()
        self._llm = llm
        log.info("loader.docx.initialised", llm_enrichment=llm is not None)

    def supported_mime_types(self) -> list[str]:
        return [
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ]

    async def load(self, raw_doc: RawDocument) -> ParsedDocument:
        # ── Stage 1: Inspect document complexity ─────────────────────────────
        complexity = await asyncio.to_thread(self._inspect_complexity, raw_doc.content)

        log.debug(
            "loader.docx.complexity",
            file=raw_doc.file_name,
            **complexity,
        )

        # ── Stage 2: Parse with MarkItDown ────────────────────────────────────
        def _convert_sync() -> str:
            stream = BytesIO(raw_doc.content)
            result = self._converter.convert_stream(
                stream,
                file_extension="docx",
            )
            return result.text_content or ""

        try:
            markdown_text = await asyncio.to_thread(_convert_sync)
        except Exception as e:
            raise DocumentParsingError(
                f"MarkItDown failed to parse '{raw_doc.file_name}': {e}"
            ) from e

        if not markdown_text.strip():
            raise DocumentParsingError(
                f"MarkItDown returned empty content for '{raw_doc.file_name}'"
            )

        # ── Stage 3: LLM enrichment — only if needed ─────────────────────────
        enrichment_applied = []

        if self._llm is not None:

            if complexity["has_images"]:
                markdown_text = await self._enrich_images(
                    raw_doc, markdown_text, complexity["image_count"]
                )
                enrichment_applied.append("image_descriptions")

            if complexity["has_nested_tables"]:
                markdown_text = await self._enrich_nested_tables(raw_doc, markdown_text)
                enrichment_applied.append("nested_table_cleanup")

        elif complexity["has_images"] or complexity["has_nested_tables"]:
            # No LLM available — log warning so operator knows quality may be reduced
            log.warning(
                "loader.docx.llm_not_configured",
                file=raw_doc.file_name,
                has_images=complexity["has_images"],
                has_nested_tables=complexity["has_nested_tables"],
                message="Complex document detected but LLM not configured — "
                "images skipped, nested tables may be incomplete",
            )

        log.info(
            "loader.docx.parsed",
            file=raw_doc.file_name,
            enrichment=enrichment_applied or "none",
            **complexity,
        )

        return ParsedDocument(
            markdown_text=markdown_text,
            extracted_metadata={
                "loader": "markitdown",
                "llm_enrichment": enrichment_applied,
                **complexity,
            },
        )

    # ── Complexity detection ──────────────────────────────────────────────────

    def _inspect_complexity(self, content: bytes) -> dict:
        """
        Inspect DOCX structure using python-docx BEFORE parsing.
        Fast — no conversion, just XML inspection.
        Returns dict with complexity flags.
        """
        from docx import Document

        doc = Document(BytesIO(content))

        # Count inline images (pictures embedded in the document body)
        image_count = len(doc.inline_shapes)

        # Detect nested tables — MarkItDown silently drops inner table content
        has_nested_tables = False
        table_count = len(doc.tables)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.tables:
                        has_nested_tables = True
                        break
                if has_nested_tables:
                    break
            if has_nested_tables:
                break

        return {
            "has_images": image_count > 0,
            "image_count": image_count,
            "has_nested_tables": has_nested_tables,
            "table_count": table_count,
            "needs_llm": image_count > 0 or has_nested_tables,
        }

    # ── LLM enrichment ────────────────────────────────────────────────────────

    async def _enrich_images(
        self,
        raw_doc: RawDocument,
        markdown_text: str,
        image_count: int,
    ) -> str:
        """
        Extract images from DOCX and add GPT-4o Vision descriptions
        inline into the Markdown text.
        """
        import base64
        from docx import Document
        from langchain_core.messages import HumanMessage

        log.info(
            "loader.docx.enriching_images",
            file=raw_doc.file_name,
            image_count=image_count,
        )

        doc = Document(BytesIO(raw_doc.content))
        descriptions = []

        for i, shape in enumerate(doc.inline_shapes):
            try:
                # Extract image bytes from the DOCX relationship
                image_part = shape._inline.graphic.graphicData.pic.blipFill.blip
                rId = image_part.get(
                    "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                )
                if not rId:
                    continue

                image_bytes = doc.part.related_parts[rId].blob
                image_b64 = base64.b64encode(image_bytes).decode()

                # Detect image format
                fmt = "jpeg"
                if image_bytes[:8] == b"\x89PNG\r\n\x1a\n":
                    fmt = "png"
                elif image_bytes[:4] == b"GIF8":
                    fmt = "gif"

                # Call Vision via your LangChain wrapper
                message = HumanMessage(
                    content=[
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/{fmt};base64,{image_b64}",
                                "detail": "high",
                            },
                        },
                        {
                            "type": "text",
                            "text": (
                                "Describe this image concisely for a RAG knowledge base. "
                                "Focus on: key information, data shown, any text visible, "
                                "and what the image communicates. "
                                "Format as a single paragraph. Be specific and factual."
                            ),
                        },
                    ]
                )

                response = await self._llm.ainvoke([message])
                description = response.content.strip()
                descriptions.append(f"\n\n**[Image {i+1}]**: {description}\n")

                log.debug(
                    "loader.docx.image_described",
                    file=raw_doc.file_name,
                    image_index=i + 1,
                )

            except Exception as e:
                log.warning(
                    "loader.docx.image_description_failed",
                    file=raw_doc.file_name,
                    image_index=i + 1,
                    error=str(e),
                )
                descriptions.append(
                    f"\n\n**[Image {i+1}]**: _Image could not be described_\n"
                )

        # Append all image descriptions at end of markdown
        if descriptions:
            markdown_text += "\n\n## Embedded Images\n" + "".join(descriptions)

        return markdown_text

    async def _enrich_nested_tables(
        self,
        raw_doc: RawDocument,
        markdown_text: str,
    ) -> str:
        """
        Use LLM to clean up nested table content that MarkItDown
        may have dropped or garbled.
        """
        from langchain_core.messages import HumanMessage, SystemMessage

        log.info(
            "loader.docx.enriching_nested_tables",
            file=raw_doc.file_name,
        )

        messages = [
            SystemMessage(
                content=(
                    "You are a document formatting assistant. "
                    "You receive Markdown converted from a Word document that may have "
                    "incomplete or garbled table content due to nested tables. "
                    "Fix the tables to be clean, complete Markdown. "
                    "Preserve all text content. Do not add or remove information. "
                    "Return only the corrected Markdown, no explanation."
                )
            ),
            HumanMessage(content=markdown_text),
        ]

        try:
            response = await self._llm.ainvoke(messages)
            cleaned = response.content.strip()
            if cleaned:
                log.debug(
                    "loader.docx.nested_tables_cleaned",
                    file=raw_doc.file_name,
                )
                return cleaned
        except Exception as e:
            log.warning(
                "loader.docx.nested_table_cleanup_failed",
                file=raw_doc.file_name,
                error=str(e),
            )

        return markdown_text  # fall back to original if LLM fails
